from docx import Document


def generate_docx(text):

    if not text:
        text = "No content available."

    # Create new document
    doc = Document()

    # Add title
    doc.add_heading("ATS Optimized Resume", 0)

    # Add resume content
    doc.add_paragraph(text)

    # Save file
    file_path = "enhanced_resume.docx"

    doc.save(file_path)

    return file_path
